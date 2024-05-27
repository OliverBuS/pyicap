from abc import ABC, abstractmethod
from io import BytesIO
# For extracting text
from pdfminer.high_level import extract_text, extract_pages
from pdfminer.layout import LTTextContainer, LTChar, LTTextLine
from docx import Document
import fitz  # PyMuPDF library
from typing import Dict

class AnalysisResult:
    def __init__(self, censor_dict: Dict[str, str], block: bool):
        self.censor_dict = censor_dict
        self.block = block

class FileOperations(ABC):
    def __init__(self, analyze_function) -> None:
        self.analyze_function = analyze_function
    @abstractmethod
    def analyze_content(self, content, file_content) -> AnalysisResult:
        pass
    @abstractmethod
    def modify_content(self, content, file_content, censor_dict: Dict[str, str]):
        pass

class TextOperations(FileOperations):
    def analyze_content(self, content, file_content) -> AnalysisResult:
        return self.analyze_function(content.decode("utf-8"))
    
    def modify_content(self, content, file_content, censor_dict: Dict[str, str]):
        print(censor_dict)
        modified_text = content.decode("utf-8")
        for key in censor_dict.keys():
            modified_text = modified_text.replace(key, censor_dict[key])
        return modified_text.encode("utf-8")

class DOCOperations(FileOperations):
    def analyze_content(self, content, file_content) -> AnalysisResult:
        # Create a BytesIO object from the file content
        docx_buffer = BytesIO(file_content)

        # Load the Word document
        document = Document(docx_buffer)

        text = ""
        # Iterate over each paragraph
        for paragraph in document.paragraphs:
            # Extract the text from the paragraph
            text += paragraph.text

        # Analyze the text using the provided function
        return self.analyze_function(text)

    def modify_content(self, content, file_content, censor_dict: Dict[str, str]):
        # Create a BytesIO object from the file content
        docx_buffer = BytesIO(file_content)

        # Load the Word document
        document = Document(docx_buffer)

        # Iterate over each paragraph
        for paragraph in document.paragraphs:
            # Extract the text from the paragraph
            text = paragraph.text
            # Modify the text using the provided function
            modified_text = text
            for key in censor_dict.keys():
                modified_text = modified_text.replace(key, censor_dict[key])
            
            print("Original text: " + text + "\nModified text: " + modified_text)

            # Replace the paragraph text with the modified text
            paragraph.text = modified_text

        # Create a BytesIO object to store the modified Word document
        output_buffer = BytesIO()
        document.save(output_buffer)

        # Get the modified Word document content as bytes
        modified_file_content = output_buffer.getvalue()

        # Reconstruct the multipart/form-data with the modified file content
        boundary = content.split(b"\r\n")[0]
        modified_content = b""
        for part in content.split(boundary)[:-1]:
            if b'filename="' in part:
                headers, _ = part.split(b"\r\n\r\n", 1)
                modified_content += (
                    boundary
                    + headers
                    + b"\r\n\r\n"
                    + modified_file_content
                    + b"\r\n"
                )
            elif len(part) > 0:
                modified_content += part

        # Add the closing boundary to the modified content
        modified_content += boundary + b"--\r\n"
        print(f"Modified content: {modified_content[-300:]}")
        return modified_content

class PDFOperations(FileOperations):
    def analyze_content(self, content, file_content) -> AnalysisResult:
        # Create a BytesIO object from the file content
        pdf_buffer = BytesIO(file_content)
        text = ""
        for page_layout in extract_pages(pdf_buffer):
            for text_container in page_layout:
                if isinstance(text_container, LTTextContainer):

                    # The element is a LTTextContainer, containing a paragraph of text.
                    text += text_container.get_text()
        # Analyze the text using the provided function
        return self.analyze_function(text)

    def modify_content(self, content, file_content, censor_dict: Dict[str, str]):
        # Create a BytesIO object from the file content
        pdf_buffer = BytesIO(file_content)
        # Open the original PDF
        pdf_file = fitz.open("pdf", pdf_buffer)

        # Iterate through each page
        for page_index in range(len(pdf_file)):
            page = pdf_file[page_index]
            
            for key in censor_dict.keys():
                censor_instances = page.search_for(key)

                for inst in censor_instances:
                    area = inst.irect  
                    redact_area = fitz.Rect(area.x0, area.y0, area.x1, area.y1)  
                    page.add_redact_annot(redact_area, text=censor_dict[key], align=fitz.TEXT_ALIGN_CENTER)  

            page.apply_redactions() 
        
        modified_file_content = pdf_file.tobytes()
        
        pdf_file.close()
        # Reconstruct the multipart/form-data with the modified file content
        boundary = content.split(b"\r\n")[0]
        modified_content = b""
        for part in content.split(boundary)[:-1]:
            if b'filename="' in part:
                headers, _ = part.split(b"\r\n\r\n", 1)
                modified_content += (
                    boundary
                    + headers
                    + b"\r\n\r\n"
                    + modified_file_content
                    + b"\r\n"
                )
            elif len(part) > 0:
                modified_content += part

        # Add the closing boundary to the modified content
        modified_content += boundary + b"--\r\n"

        print(f"Modified content: {modified_content[-300:]}")
        return modified_content